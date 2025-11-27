#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Générateur de cartes d'oracles (D&D / JDR Fantastique).

- Lit la configuration depuis un fichier JSON.
- Demande à l'utilisateur combien de cartes générer.
- Vérifie la validité des listes critiques dans la configuration (Robustesse A1).
- Utilise la distribution de titres définie dans "title_distribution".
- **Produit un fichier texte avec un format encadré (Nouveau)**.
- Optionnel : produit un DOCX si python-docx est installé.

Usage :
    python generate_deck.py [chemin/vers/deck_config.json]
"""

import json
import random
import sys
from pathlib import Path
from itertools import chain
from typing import Any, List, Dict, Union

# -----------------------------------------------------
#  Utilitaire DOCX optionnel
# -----------------------------------------------------
try:
    from docx import Document   # Nécessite : pip install python-docx
    from docx.shared import Inches
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# -----------------------------------------------------
#  Fonctions d'aide et de configuration
# -----------------------------------------------------

def load_config(config_path: Path) -> dict:
    """Charge le fichier JSON de configuration (UTF-8)."""
    if not config_path.is_file():
        raise FileNotFoundError(f"Fichier de configuration introuvable : {config_path}")
    try:
        with open(config_path, "r", encoding="utf-8") as f:
            return json.load(f)
    except json.JSONDecodeError as e:
        raise ValueError(f"Erreur de décodage JSON dans {config_path}: {e}")


def build_title_pool(title_distribution: dict, card_count: int) -> list:
    """Construit une liste de titres de longueur 'card_count' selon la distribution."""
    items = list(title_distribution.items())
    counts = [int(v) for _, v in items]
    total = sum(counts)

    if total >= card_count:
        pool = chain.from_iterable([title] * count for (title, _), count in zip(items, counts))
        pool = list(pool)
        random.shuffle(pool)
        return pool[:card_count]
    else:
        titles = [t for t, _ in items]
        weights = [float(v) for _, v in items]
        return random.choices(titles, weights=weights, k=card_count)


def pick_multiple(source_list: List[Any], count: int) -> List[Any]:
    """Renvoie une liste de 'count' éléments distincts ou avec répétitions."""
    if not source_list:
        return []
    if len(source_list) < count:
        return [random.choice(source_list) for _ in range(count)]
    return random.sample(source_list, count)


def _get_random_field_value(cfg: dict, field_key: str, count: int = 1) -> Union[str, List[str]]:
    """Récupère un champ depuis la config, gère la sélection simple/multiple."""
    source_list = cfg.get(field_key)
    
    if not source_list:
        return "" if count == 1 else []

    if count == 1:
        return random.choice(source_list)
    else:
        return pick_multiple(source_list, count)


# -----------------------------------------------------
#  Génération de carte
# -----------------------------------------------------

def generate_card(index: int, title: str, cfg: dict) -> Dict[str, Any]:
    """
    Génère une carte sous forme de dict, utilisant uniquement les clés du JSON.
    """
    card = {
        "number": index,
        "title": title,
        "symbol": _get_random_field_value(cfg, "symbols"),
        "verbs": _get_random_field_value(cfg, "table_verbes", 3), 
        "lieu": _get_random_field_value(cfg, "lieux"),
        "personnage": _get_random_field_value(cfg, "personnages"),
        "objet": _get_random_field_value(cfg, "objets"),
        "emotions": _get_random_field_value(cfg, "emotions", 2),
        "appearance": _get_random_field_value(cfg, "appearances"),
        "motivation": _get_random_field_value(cfg, "motivations"),
        "traits": _get_random_field_value(cfg, "traits", 3),
        "secret": _get_random_field_value(cfg, "sombres_secrets"), 
        "reaction": _get_random_field_value(cfg, "reactions_amical_hostile"),
        "relation": _get_random_field_value(cfg, "relations_pj_pnj"), 
        "borders": _get_random_field_value(cfg, "borders", min(12, len(cfg.get("borders", [])) or 0))
    }

    return {k: v for k, v in card.items() if v}


# -----------------------------------------------------
#  Fonctions d'enregistrement et d'affichage
# -----------------------------------------------------

def format_card_as_text(card: dict) -> str:
    """
    **Retourne une chaîne de caractères formatée en boîte encadrée d'astérisques.**
    """
    # 1. Préparation des lignes de contenu (format "Label : Value")
    content_lines = []
    
    # Titre principal
    content_lines.append(f"Carte {card.get('number')} — {card.get('title', '')}")

    fields = [
        ("symbol", "Symbole", ""), ("verbs", "Verbes", ", "), ("lieu", "Lieu", ""), 
        ("personnage", "Personnage", ""), ("objet", "Objet", ""), ("emotions", "Émotions", ", "), 
        ("appearance", "Apparence", ""), ("motivation", "Motivation", ""), 
        ("traits", "Traits", ", "), ("secret", "Secret", ""), ("relation", "Relation", "")
    ]
    
    for key, label, separator in fields:
        value = card.get(key)
        if value:
            if isinstance(value, list):
                content_lines.append(f"{label} : {separator.join(value)}")
            else:
                content_lines.append(f"{label} : {value}")

    if card.get("reaction"):
        # Le format d'affichage doit correspondre exactement à celui de l'utilisateur
        content_lines.append(f"Réaction (amical/hostile) : {card['reaction']}")

    if card.get("borders"):
        content_lines.append(f"Thèmes : {', '.join(card['borders'])}")
    
    if not content_lines:
        return ""

    # 2. Détermination de la longueur maximale et du cadre
    max_len = max(len(line) for line in content_lines)
    
    # Largeur de la boîte : longueur max + 2 espaces de chaque côté + 2 astérisques (total + 4)
    box_width = max_len + 4 
    
    top_bottom_line = "*" * box_width
    
    # 3. Construction de la carte encadrée
    boxed_lines = [top_bottom_line]
    
    for line in content_lines:
        # Calcul du padding pour aligner à droite
        # On utilise un espace de padding de 2, donc (box_width - len(line) - 2 astérisques) / 2
        # Pour une boîte parfaite, on fait : (longueur totale de la ligne - longueur du contenu - 2)
        padding = " " * (max_len - len(line))
        # Format: *Ligne de contenu + padding*
        boxed_lines.append(f"*{line}{padding} *")

    boxed_lines.append(top_bottom_line)
    
    # Ajout d'une ligne vide à la fin pour la séparation entre les cartes
    return "\n".join(boxed_lines) + "\n"


def _find_symbol_image(symbol: str):
    """Recherche d'images pour le docx (inchangée)."""
    candidates = []
    for base in (Path.cwd() / "symbols", Path.cwd()):
        for ext in (".jpg", ".png", ".jpeg", ".webp"):
            candidates.append(base / f"{symbol}{ext}")
    for p in candidates:
        if p.is_file():
            return p
    return None


def save_as_txt(cards: List[dict], output_path: str):
    """Enregistre toutes les cartes dans un fichier texte."""
    path = Path(output_path)
    content = []
    deck_title = "LE TAROT DES ROYAUMES OUBLIÉS — DECK GÉNÉRÉ"
    content.append(deck_title + "\n")
    content.append("=" * len(deck_title) + "\n")

    for card in cards:
        # Utilisation de la fonction formatée
        content.append(format_card_as_text(card))

    path.write_text("\n".join(content), encoding="utf-8")
    print(f"[OK] Fichier texte généré : {path}")


def save_as_docx(cards: List[dict], output_path: str):
    """Enregistre toutes les cartes dans un fichier DOCX."""
    if not HAS_DOCX:
        print("[INFO] python-docx n’est pas installé, DOCX non généré.")
        return

    doc = Document()
    
    for card in cards:
        doc.add_heading(f"{card['number']} :  {card['title']}", level=2)

        # Logique d'ajout du Symbole
        symbol = card.get("symbol", "")
        img_path = _find_symbol_image(symbol) if symbol else None
        if img_path is not None:
            p = doc.add_paragraph()
            r = p.add_run("")
            r.add_picture(str(img_path), width=Inches(0.52))
        else:
            doc.add_paragraph(f"Symbole : {symbol}")

        # Ajout des champs (simplifiés)
        if card.get("verbs"):
            doc.add_paragraph(f"Action(s) centrale(s) : {', '.join(card['verbs'])}")
        if card.get("lieu"):
            doc.add_paragraph(f"Lieu : {card['lieu']}")
        if card.get("personnage"):
            doc.add_paragraph(f"Personnage : {card['personnage']}")
        if card.get("objet"):
            doc.add_paragraph(f"Objet : {card['objet']}")
        if card.get("motivation"):
            doc.add_paragraph(f"Motivation : {card['motivation']}")
        if card.get("traits"):
            doc.add_paragraph(f"Caractère : {', '.join(card['traits'])}")
        if card.get("secret"):
            doc.add_paragraph(f"Secret : {card['secret']}")
        if card.get("relation"):
            doc.add_paragraph(f"Relation : {card['relation']}")

        if card.get("reaction"):
            doc.add_paragraph(f"Réaction (amical/hostile) : {card['reaction']}")

        if card.get("borders"):
            doc.add_paragraph(f"Thèmes dominants : {', '.join(card['borders'])}")

        doc.add_paragraph("")

    path = Path(output_path)
    doc.save(path)
    print(f"[OK] Fichier DOCX généré : {path}")


def _prompt_for_card_count(default_count: int) -> int:
    """Demande à l'utilisateur le nombre de cartes à générer."""
    while True:
        try:
            prompt = f"Combien de cartes générer ? (Défaut: {default_count}) : "
            user_input = input(prompt).strip()
            
            if not user_input:
                return default_count
            
            count = int(user_input)
            if count <= 0:
                print("[ATTENTION] Le nombre doit être supérieur à zéro.")
                continue
            return count
        except ValueError:
            print("[ATTENTION] Entrée invalide. Veuillez entrer un nombre entier positif.")


def check_critical_lists(cfg: dict):
    """Vérifie que les listes essentielles existent et ne sont pas vides."""
    required_keys = [
        "title_distribution", "symbols", "table_verbes", "lieux", 
        "personnages", "objets", "motivations", "traits", 
        "sombres_secrets", "reactions_amical_hostile", "relations_pj_pnj"
    ]
    
    missing_keys = []
    empty_keys = []
    
    for key in required_keys:
        if key not in cfg:
            missing_keys.append(key)
        elif not cfg.get(key) or (isinstance(cfg.get(key), (list, dict)) and not cfg.get(key)):
            empty_keys.append(key)

    if missing_keys or empty_keys:
        error_msg = "[ERREUR FATALE] La configuration JSON est incomplète ou vide :\n"
        if missing_keys:
            error_msg += f"- Clés **manquantes** : {', '.join(missing_keys)}\n"
        if empty_keys:
            error_msg += f"- Clés **vides** (doivent contenir des éléments) : {', '.join(empty_keys)}\n"
        error_msg += "Veuillez vérifier votre fichier deck_config.json."
        raise ValueError(error_msg)


def main():
    # Détermination du chemin de la config (robuste)
    if len(sys.argv) > 1:
        config_path = Path(sys.argv[1]).resolve()
    else:
        config_path = Path("deck_config.json").resolve()
    
    try:
        cfg = load_config(config_path)

        # Vérification de la configuration critique
        check_critical_lists(cfg)

        # Demander le nombre de cartes
        default_count = int(cfg.get("card_count", 100))
        card_count = _prompt_for_card_count(default_count)

        # Construction du pool de titres
        title_pool = build_title_pool(cfg["title_distribution"], card_count)

    except (FileNotFoundError, ValueError) as e:
        print(f"\n[ERREUR] Impossible de lancer le générateur :\n{e}")
        sys.exit(1)

    # Génération des cartes
    cards = []
    for i in range(1, card_count + 1):
        title = title_pool[i - 1]
        card = generate_card(i, title, cfg)
        cards.append(card)

    # Sorties
    out_cfg = cfg.get("output", {})
    txt_file = out_cfg.get("txt_file", "deck_oracle.txt")
    docx_file = out_cfg.get("docx_file", "deck_oracle.docx")
    create_docx = bool(out_cfg.get("create_docx", True))

    save_as_txt(cards, txt_file)

    if create_docx:
        save_as_docx(cards, docx_file)


if __name__ == "__main__":
    main()